from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path.cwd() / "冒险书产品说明书.docx"

INK = RGBColor(42, 33, 25)
MUTED = RGBColor(92, 78, 61)
BLUE = RGBColor(46, 116, 181)
LIGHT = "F4F0E6"
HEADER = "E8EEF5"


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=11, bold=False, color=INK, after=6, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 7)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, color=BLUE, bold=True)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1900, 7460])
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        left, right = table.rows[index].cells
        if index == 0:
            set_cell_shading(left, HEADER)
            set_cell_shading(right, HEADER)
        else:
            set_cell_shading(left, LIGHT)
        left.text = ""
        right.text = ""
        set_run_font(left.paragraphs[0].add_run(label), size=10, color=MUTED, bold=True)
        set_run_font(right.paragraphs[0].add_run(value), size=10.5, color=INK)
    add_para(doc, "", after=4)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, widths)
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        set_cell_shading(cell, HEADER)
        cell.text = ""
        set_run_font(cell.paragraphs[0].add_run(header), size=10, color=INK, bold=True)
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            cell = table.rows[row_index].cells[col]
            cell.text = ""
            set_run_font(cell.paragraphs[0].add_run(value), size=10, color=INK)
    add_para(doc, "", after=4)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    header = section.header.paragraphs[0]
    header.text = ""
    set_run_font(header.add_run("冒险书 | 产品说明书"), size=9, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    set_run_font(footer.add_run("用于统一产品理解、后续设计讨论与功能规划"), size=9, color=MUTED)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "产品说明书", size=10, bold=True, color=MUTED, after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("冒险书"), size=28, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("一款带有游戏感的个人事务管理工具"), size=13, color=MUTED)

    add_label_table(
        doc,
        [
            ("文档用途", "说明当前产品的整体结构、核心功能、页面分工与后续扩展方向"),
            ("说明范围", "只描述产品与使用体验，不展开具体实现方式"),
            ("当前阶段", "概念与原型逐步成型阶段，任务体系和 AI 向导正在细化"),
        ],
    )

    add_heading(doc, "1. 产品定位")
    add_para(
        doc,
        "冒险书是一款把个人事务管理包装成“冒险记录”的工具。它不是单纯的待办清单，而是希望让用户把长期目标、重复习惯、一次性事务、日历安排和日常记录放进同一本“冒险手册”里。",
    )
    add_para(
        doc,
        "产品的核心感受应当是：安静、复古、像纸质档案或任务契约，同时又有游戏里的任务书、日常任务、支线任务和旅行向导。",
    )
    for text in [
        "用“今日任务”帮助用户知道现在该做什么。",
        "用“任务书”承载有起止日期、任务线和子任务的长期事项。",
        "用“日常任务”承载每日、每周或自定义周期的重复行动。",
        "用“支线任务”承载完成一次就归档的独立事务。",
        "用“时间线”和“冒险笔记”帮助用户回看安排、记录当天状态。",
        "用“旅行向导安雅”等角色，把任务梳理和任务创建变成对话。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. 产品整体结构")
    add_para(
        doc,
        "冒险书的主体验由一个今日主界面和若干子界面组成。用户可以从侧边导航或底部导航切换不同区域；在移动视图中，导航应保持稳定、清晰，不因为页面变长而改变样式。",
    )
    add_matrix(
        doc,
        ["区域", "主要作用", "用户会在这里完成什么"],
        [
            ("今日", "当前行动入口", "查看今日截止、已逾期、即将到来的任务；查看状态仪表；写今日便签；呼出安雅。"),
            ("任务书", "长期线索管理", "建立一本任务书，设置起止日期，安排多条任务线、子任务和独立任务。"),
            ("日常任务", "重复行动管理", "创建每日、每周或自定义周期的任务，并在当前周期内重复完成。"),
            ("支线任务", "一次性事务管理", "记录独立任务，完成后归档，适合零散但明确的事项。"),
            ("时间线", "日期视角", "通过日历查看某一天相关任务，并维护当天备注。"),
            ("冒险笔记", "随笔与日志", "记录想法、备忘、复盘和与任务无直接绑定的内容。"),
            ("设置", "偏好与角色能力", "切换语言，后续接入 AI 向导所需配置。"),
        ],
        [1400, 2200, 5760],
    )

    add_heading(doc, "3. 今日主界面")
    add_para(
        doc,
        "今日是用户打开产品后最重要的页面。它不应是复杂的大控制台，而应该像一张“今日清单”：用户一眼看到今天要处理什么、哪些已经拖延、哪些马上临近。",
    )
    for text in [
        "顶部保留清晰标题和新增入口。",
        "状态仪表显示任务总数、待执行、已完成和完成进度。",
        "AI Agent 区域放置旅行向导安雅的头像，点击后进入对话。",
        "今日便签用于记录当天备注，来源与时间线中的当天备注保持一致。",
        "任务清单使用档案夹标签切换“全部、任务书、日常、支线”等视角。",
        "清单内部更适合用淡色分隔线，而不是每条任务都做成厚重卡片。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4. 三种任务类型")
    add_matrix(
        doc,
        ["任务类型", "定位", "关键规则"],
        [
            ("任务书", "长期目标、项目、连续事件", "先设置起止日期；内部可有多条并行任务线和独立任务；任务按时间顺序排列。"),
            ("日常任务", "重复习惯和周期事项", "设置有效期间和重复周期；在周期内可以多次完成；适合打卡和例行事项。"),
            ("支线任务", "一次性独立事项", "完成一次后进入归档；不需要复杂结构；适合临时委托和单点事务。"),
        ],
        [1600, 2600, 5160],
    )

    add_heading(doc, "5. 任务书系统")
    add_para(doc, "任务书是当前产品最具特色的部分。每一本任务书都像一个小型战役或章节，它有自己的起止日期、整体进度、任务线和独立任务。")
    add_heading(doc, "任务书的内部结构", 2)
    for text in [
        "先创建一本任务书，并设置名称、开始日期和结束日期。",
        "在任务书中添加一条或多条任务线。任务线可以理解为同一个目标下连续推进的一串子任务。",
        "每条任务线下可以添加多个子任务，每个子任务独立设置时间。",
        "任务书中也可以放置独立任务，它们不属于某条任务线，但仍参与整本任务书的时间安排。",
        "任务线和独立任务会在横向时间轴上按日期排列，形成类似战役排程图的效果。",
    ]:
        add_number(doc, text)
    add_heading(doc, "任务书页面应呈现的感觉", 2)
    for text in [
        "顶部是任务书标题、进度和管理入口。",
        "中间是横向时间轴，能看到起点、终点和关键日期。",
        "时间轴下方是多条并行任务线，每条任务线像一条路线。",
        "子任务使用轻量单线框，不使用厚重双重框。",
        "任务跨度应能通过卡片长度表现出来，而不是只作为一个日期点。",
        "连续没有任务的空白时间可以适度压缩，用省略段表示，以减少横向空间浪费。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "6. 日常任务与支线任务")
    add_para(doc, "日常任务和支线任务都应保留档案夹式标签设计，用于在“全部、进行中、归档”等状态间切换。每个子界面只管理自己的任务类型，不应出现一个全局任务编辑栏。")
    add_heading(doc, "日常任务", 2)
    for text in ["适合每日复盘、运动、喝水、阅读、每周整理等重复事项。", "需要有效期间、重复周期、当前完成状态和连续次数等信息。", "完成后不一定立刻消失，而是进入对应周期的完成状态。"]:
        add_bullet(doc, text)
    add_heading(doc, "支线任务", 2)
    for text in ["适合“买某件东西”“联系某人”“提交某个材料”等单次事务。", "完成一次后进入归档。", "页面应比任务书更轻，不需要任务线或复杂排程。"]:
        add_bullet(doc, text)

    add_heading(doc, "7. 时间线与便签")
    add_para(doc, "时间线提供日历视角。用户可以选择某一天，查看当天相关的任务书子任务、日常任务和支线任务。选中日期下方的任务查看逻辑应与今日页类似，支持按档案夹切换不同类型。")
    for text in ["日历和任务查看区应上下排列，而不是左右挤在一起。", "选中某天后，显示当天任务和当天备注。", "今日页的便签与时间线中当天备注应保持同一份内容。", "时间线不是复杂分析面板，而是帮助用户回看与安排日期。"]:
        add_bullet(doc, text)

    add_heading(doc, "8. 冒险笔记")
    add_para(doc, "冒险笔记是随笔、备忘录和日志区域。它补足了任务系统无法表达的内容：想法、灵感、心情、复盘、计划草稿、会议记录等。")
    for text in ["可以创建普通笔记。", "可以作为日常日志使用。", "未来可以与任务、日期、角色对话产生关联。", "视觉上应保持纸张和档案感，不变成普通富文本编辑器。"]:
        add_bullet(doc, text)

    add_heading(doc, "9. AI 向导与角色体系")
    add_para(doc, "第一位 AI 向导是“旅行向导 安雅”。她不是一个普通聊天窗口，而是产品世界观中的角色。用户在今日页点击她的头像，可以展开居中的游戏式对话框。")
    add_heading(doc, "安雅当前承担的能力", 2)
    for text in [
        "读取当前任务清单，告诉用户今天有哪些任务。",
        "提醒已逾期和临近截止的任务。",
        "根据任务重要性、截止日期和任务类型，建议今天的处理顺序。",
        "通过多轮对话理解用户描述，并帮助创建任务书、日常任务或支线任务。",
        "在真正创建任务前，必须先向用户确认。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "后续角色扩展", 2)
    add_para(doc, "未来可以加入更多 AI Agent，每个角色负责不同方向。例如复盘助手、习惯教练、项目参谋、写作记录员等。角色应分工明确，不把所有能力堆在同一个对话入口里。")

    add_heading(doc, "10. 视觉与交互原则")
    for text in [
        "保持复古纸张、档案夹、契约、任务书的整体风格。",
        "不要随意替换现有 UI 语言和组件气质。",
        "主界面保持简单，重点是今日清单，不做复杂控制面板。",
        "档案夹标签是重要设计元素，不能随意删除。",
        "移动端底部导航应稳定、清楚、始终保持一致。",
        "任务列表内部尽量用细线和淡色分隔，避免到处都是厚重卡片。",
        "图标按钮可以纯图案表达，不必每个按钮都有边框和文字。",
        "角色头像、便签、状态仪表等模块应像纸面元素一样自然存在。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "11. 典型使用流程")
    add_heading(doc, "早上打开冒险书", 2)
    for text in ["进入今日页。", "查看状态仪表，知道今天任务规模。", "点击安雅头像，让她梳理今日路线。", "根据建议先处理逾期或高优先级任务。", "在今日便签记录今天要记住的事。"]:
        add_number(doc, text)
    add_heading(doc, "创建一本长期任务书", 2)
    for text in ["进入任务书页面。", "创建任务书，填写名称、开始日期和结束日期。", "添加多条任务线，并为每条任务线添加子任务。", "为每个子任务设置时间。", "在横向时间轴中查看整个任务书的推进节奏。"]:
        add_number(doc, text)
    add_heading(doc, "通过安雅创建任务", 2)
    for text in ["打开安雅对话。", "用自然语言说明想做什么。", "安雅判断适合创建任务书、日常任务还是支线任务。", "安雅给出创建草案。", "用户确认后，任务才正式加入。"]:
        add_number(doc, text)

    add_heading(doc, "12. 后续规划建议")
    add_matrix(
        doc,
        ["方向", "说明"],
        [
            ("任务书编辑体验", "继续优化任务书的起止日期、任务线、子任务和独立任务管理，让它像真正的战役规划图。"),
            ("AI 角色扩展", "在安雅之外加入更多角色，让不同角色处理不同类型的帮助。"),
            ("笔记联动", "让冒险笔记可以关联日期、任务书或某次对话。"),
            ("更细的归档体系", "让完成的任务、笔记和任务书有更清晰的回看方式。"),
            ("个性化设置", "除了语言切换，逐步加入角色配置、显示偏好和数据导入导出等能力。"),
        ],
        [2200, 7160],
    )

    add_heading(doc, "13. 一句话总结")
    add_para(
        doc,
        "冒险书要成为一套“用游戏任务感组织现实生活”的个人事务系统：今天做什么由今日页负责，长期目标由任务书负责，重复行动由日常任务负责，零散事务由支线任务负责，时间线和笔记负责记录，AI 角色负责陪用户整理、判断和创建。",
        size=12,
        bold=True,
        after=0,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
