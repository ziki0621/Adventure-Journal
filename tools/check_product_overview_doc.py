from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = Path.cwd() / "冒险书产品说明书.docx"
doc = Document(path)
texts = "\n".join(paragraph.text for paragraph in doc.paragraphs)
checks = ["冒险书", "产品定位", "任务书系统", "旅行向导 安雅"]
with ZipFile(path) as archive:
    xml = archive.read("word/document.xml").decode("utf-8")

print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))
print("checks", all(item in texts for item in checks))
print("xml_has_cn", all(item in xml for item in checks[:3]))
print("question_marks", xml.count("????"))
